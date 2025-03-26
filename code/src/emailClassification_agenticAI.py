import os
from langchain import PromptTemplate, LLMChain
import google.generativeai as genai
from PyPDF2 import PdfFileReader
import email
from email import policy
from email.parser import BytesParser
from io import BytesIO
import streamlit as st
from docx import Document
from langchain_google_genai import GoogleGenerativeAI
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
import uuid

from io import BytesIO
from dotenv import load_dotenv

load_dotenv()
google_token = os.getenv("GOOGLE_API_KEY")

# Initialize Gemini Model
gemini_model = GoogleGenerativeAI(model="gemini-1.5-pro", google_api_key=google_token)

# Define Prompt & Chain
classifier_chain = LLMChain(
    prompt=PromptTemplate(
        input_variables=["content","subject"],
        template="""
        Classify the following email subject and content into a service request type.
        Use these rules for classification:
        - Keywords like 'Transfer', 'Fund', or 'Wire' indicate AU Transfer.
        - Words like 'Close', 'Termination', or 'Final Notice' suggest Closing Notice.
        - Mentions of 'Fee', 'Payment', or 'Charge' indicate Fee Payment.
        - References to 'Balance', 'Due', or 'Outstanding' suggest Outstanding Balance.
        - General inquiries or unrelated content should be labeled as General Inquiry.
        
        Subject: {subject}
        Content: {content}
        """
    ),
    llm=gemini_model
)


# Content Extraction Agent
def extract_text_from_file(file):
    content = ""
    if file.name.endswith(".docx"):
        doc = Document(file)
        content = "\n".join([para.text for para in doc.paragraphs])
    elif file.name.endswith(".pdf"):
        reader = PdfFileReader(file)
        for page in reader.pages:
            content += page.extract_text() + "\n"
    elif file.name.endswith(".eml"):
        msg = BytesParser(policy=policy.default).parse(file)
        body_content = msg.get_body(preferencelist=("plain",)).get_content() if msg.get_body() else ""
        content += body_content

        # Extract attachments text if body content is insufficient
        if len(content.strip()) < 100:  # Prioritize body if sufficient
            for part in msg.iter_attachments():
                if part.get_filename():
                    attachment_data = part.get_payload(decode=True)
                    attachment_file = BytesIO(attachment_data)
                    try:
                        if part.get_filename().endswith(".docx"):
                            doc = Document(attachment_file)
                            content += "\n" + "\n".join([para.text for para in doc.paragraphs])
                        elif part.get_filename().endswith(".pdf"):
                            reader = PdfFileReader(attachment_file)
                            for page_num in range(reader.numPages):
                                content += "\n" + reader.getPage(page_num).extract_text()
                    except Exception as e:
                        content += f"\n[Attachment {part.get_filename()} could not be read]"
    return content.strip()

previous_emails = set()

def classify_email(content, subject, previous_emails):
    is_duplicate = content in previous_emails

    if is_duplicate:
        return {
            "Service Request Type": "Duplicate",
            "Sub Request Type": "N/A",
            "Priority": "Low",
            "Confidence Score": 1.0,
            "Duplicate Indicator": True
        }

      # Use Gemini classifier_chain for enhanced request type identification
    result = classifier_chain.run({
    "content": content,
    "subject": subject
    })

    # Extract identified type using improved logic
    for keyword in ['AU Transfer', 'Closing Notice', 'Fee Payment', 'Outstanding Balance', 'General Inquiry']:
        if keyword.lower() in result.lower():
            identified_type = keyword
            break

    priority_map = {
        "AU Transfer": "High",
        "Closing Notice": "Medium",
        "Fee Payment": "Low",
        "Outstanding Balance": "Medium",
        "General Inquiry": "Low"
    }

    priority = priority_map.get(identified_type, "Low")
    confidence_score = 0.95  # Mock confidence for demonstration

    return {
        "Service Request Type": identified_type,
        "Sub Request Type": "Detailed Request" if identified_type != "Unknown" else "General Inquiry",
        "Priority": priority,
        "Confidence Score": confidence_score,
        "Duplicate Indicator": False
    }
# Streamlit UI
st.title("📧 Agentic AI - Email Classification using GEMINI Model")
uploaded_file = st.file_uploader("Upload a .docx, .pdf, or .eml file", type=["docx", "pdf", "eml"])


if uploaded_file is not None:
    content = extract_text_from_file(uploaded_file)
    subject = "Sample Subject"  # Modify this to extract the actual subject if needed
    result = classify_email(content, subject, previous_emails)
    previous_emails.add(content)
    st.json({"File Name": uploaded_file.name, "Classification": result})




